from __future__ import annotations

from io import BytesIO
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BLUE = "0798CF"
GREEN = "009B4C"
TEAL = "16AD8F"
PURPLE = "A990C7"
ORANGE = "F68B08"
GRAY = "858E93"
INK = "35434B"
PDF_FONT = "Helvetica"
PDF_FONT_BOLD = "Helvetica-Bold"
for regular, bold in [
    ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
]:
    if Path(regular).exists() and Path(bold).exists():
        pdfmetrics.registerFont(TTFont("CoinvierteSans", regular))
        pdfmetrics.registerFont(TTFont("CoinvierteSansBold", bold))
        PDF_FONT, PDF_FONT_BOLD = "CoinvierteSans", "CoinvierteSansBold"
        break


def _image_bytes(image) -> bytes:
    if isinstance(image, bytes):
        return image
    if hasattr(image, "getvalue"):
        return image.getvalue()
    return bytes(image)


def _usable_images(images: list) -> list[bytes]:
    valid = []
    for image in images or []:
        try:
            raw = _image_bytes(image)
            PILImage.open(BytesIO(raw)).verify()
            valid.append(raw)
        except Exception:
            continue
    return valid


def _money(value) -> str:
    return f"${float(value or 0):,.2f} MXN"


def _set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _doc_heading(doc, text: str, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _font(run, 14, True, color)
    return p


def _doc_label_value(doc, label: str, value: str):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_geometry(table, [2700, 6660])
    left, right = table.rows[0].cells
    _set_cell_fill(left, "F0F3F4")
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    _font(p.add_run(label), 9.5, True, GRAY)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    _font(p.add_run(str(value or "Sin información")), 10.5, False, INK)


def build_docx(data: dict, images: list, logo_path: str | Path) -> bytes:
    """Ficha institucional. Excluye deliberadamente Gestión Documental."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.72)
    section.left_margin = section.right_margin = Inches(0.85)
    section.header_distance = section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    if Path(logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        p.add_run().add_picture(str(logo_path), width=Inches(6.25))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run("FICHA DEL PROYECTO"), 19, True, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(13)
    _font(p.add_run(data.get("direccion", "")), 10.5, True, ORANGE)

    _doc_heading(doc, "Información General", BLUE)
    for label, value in [
        ("Nombre del proyecto", data.get("nombre")),
        ("Solicitante", data.get("solicitante")),
        ("Municipio de ejecución", data.get("municipio")),
        ("Año de inicio", data.get("anio_inicio")),
        ("Monto", _money(data.get("monto"))),
        ("Objetivo general", data.get("objetivo_general")),
    ]:
        _doc_label_value(doc, label, value)
    p = doc.add_paragraph()
    _font(p.add_run("Objetivos específicos"), 10.5, True, GRAY)
    for objective in data.get("objetivos_especificos", []):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.35)
        p.paragraph_format.first_line_indent = Inches(-.18)
        _font(p.add_run(objective), 10.5, False, INK)

    _doc_heading(doc, "Monitoreo y Seguimiento", GREEN)
    monitoring = data.get("monitoreo", {})
    for label, key, suffix in [
        ("Estatus", "estatus", ""), ("Responsable", "responsable", ""),
        ("Periodo de seguimiento", "periodo", ""), ("Avance", "avance", "%"),
        ("Principales avances", "avances", ""), ("Pendientes o riesgos", "pendientes", ""),
        ("Próximas acciones", "proximas_acciones", ""), ("Observaciones", "observaciones", ""),
    ]:
        value = monitoring.get(key, "")
        if value != "" and value is not None:
            _doc_label_value(doc, label, f"{value}{suffix}")

    valid_images = _usable_images(images)
    _doc_heading(doc, "Evidencia Fotográfica", PURPLE)
    if not valid_images:
        p = doc.add_paragraph("Sin evidencia fotográfica cargada.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, raw in enumerate(valid_images, 1):
        stream = BytesIO(raw)
        img = PILImage.open(BytesIO(raw))
        max_w, max_h = 6.15, 3.9
        ratio = min(max_w / img.width, max_h / img.height) * img.width
        width = max(.8, ratio)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_together = True
        p.add_run().add_picture(stream, width=Inches(width))
        c = doc.add_paragraph(f"Fotografía {index}")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(c.runs[0], 8.5, False, GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run("COINVIERTE · Ficha institucional del proyecto"), 8, False, GRAY)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _pdf_image(raw: bytes, max_width=6.2 * inch, max_height=4.3 * inch):
    pil = PILImage.open(BytesIO(raw))
    width, height = pil.size
    scale = min(max_width / width, max_height / height)
    return Image(BytesIO(raw), width=width * scale, height=height * scale)


def build_pdf(data: dict, images: list, logo_path: str | Path) -> bytes:
    """Ficha institucional. Excluye deliberadamente Gestión Documental."""
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter, rightMargin=.65*inch, leftMargin=.65*inch,
                            topMargin=.55*inch, bottomMargin=.55*inch, title="Ficha del proyecto")
    styles = getSampleStyleSheet()
    title = ParagraphStyle("TitleCoinvierte", parent=styles["Title"], fontName=PDF_FONT_BOLD,
                           fontSize=19, leading=22, textColor=colors.HexColor("#"+INK), alignment=TA_CENTER, spaceAfter=4)
    subtitle = ParagraphStyle("SubtitleCoinvierte", parent=styles["Normal"], fontName=PDF_FONT_BOLD,
                              fontSize=9.5, textColor=colors.HexColor("#"+ORANGE), alignment=TA_CENTER, spaceAfter=12)
    body = ParagraphStyle("BodyCoinvierte", parent=styles["BodyText"], fontName=PDF_FONT,
                          fontSize=9.2, leading=12, textColor=colors.HexColor("#"+INK))
    label = ParagraphStyle("LabelCoinvierte", parent=body, fontName=PDF_FONT_BOLD, fontSize=8.5,
                           textColor=colors.HexColor("#"+GRAY))
    story = []
    if Path(logo_path).exists():
        logo = PILImage.open(logo_path)
        scale = min(6.3*inch/logo.width, 1.25*inch/logo.height)
        story += [Image(str(logo_path), logo.width*scale, logo.height*scale), Spacer(1, 8)]
    story += [Paragraph("FICHA DEL PROYECTO", title), Paragraph(data.get("direccion", ""), subtitle)]

    def section_heading(text, color):
        table = Table([[Paragraph(text, ParagraphStyle("sh", parent=body, fontName=PDF_FONT_BOLD, fontSize=13,
                                                       textColor=colors.white))]], colWidths=[7.2*inch])
        table.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#"+color)),
                                   ("LEFTPADDING",(0,0),(-1,-1),10),("TOPPADDING",(0,0),(-1,-1),6),
                                   ("BOTTOMPADDING",(0,0),(-1,-1),6)]))
        story.extend([Spacer(1, 6), table, Spacer(1, 6)])

    def rows_table(rows):
        content = [[Paragraph(str(a), label), Paragraph(str(b or "Sin información"), body)] for a,b in rows]
        table = Table(content, colWidths=[1.75*inch,5.45*inch], repeatRows=0)
        table.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"MIDDLE"),("GRID",(0,0),(-1,-1),.35,colors.HexColor("#DDE4E6")),
                                   ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#F0F3F4")),
                                   ("LEFTPADDING",(0,0),(-1,-1),7),("RIGHTPADDING",(0,0),(-1,-1),7),
                                   ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
        story.append(table)

    section_heading("Información General", BLUE)
    rows_table([("Nombre del proyecto",data.get("nombre")),("Solicitante",data.get("solicitante")),
                ("Municipio de ejecución",data.get("municipio")),("Año de inicio",data.get("anio_inicio")),
                ("Monto",_money(data.get("monto"))),("Objetivo general",data.get("objetivo_general")),
                ("Objetivos específicos","<br/>".join(f"• {x}" for x in data.get("objetivos_especificos",[])))])
    section_heading("Monitoreo y Seguimiento", GREEN)
    m = data.get("monitoreo", {})
    rows_table([("Estatus",m.get("estatus")),("Responsable",m.get("responsable")),("Periodo",m.get("periodo")),
                ("Avance",f"{m.get('avance',0)}%"),("Principales avances",m.get("avances")),
                ("Pendientes o riesgos",m.get("pendientes")),("Próximas acciones",m.get("proximas_acciones")),
                ("Observaciones",m.get("observaciones"))])
    section_heading("Evidencia Fotográfica", PURPLE)
    valid_images = _usable_images(images)
    if not valid_images:
        story.append(Paragraph("Sin evidencia fotográfica cargada.", body))
    for idx, raw in enumerate(valid_images, 1):
        story.append(KeepTogether([_pdf_image(raw), Spacer(1,3), Paragraph(f"Fotografía {idx}",
            ParagraphStyle("caption",parent=body,fontSize=8,textColor=colors.HexColor("#"+GRAY),alignment=TA_CENTER)),Spacer(1,8)]))

    def footer(canvas, _doc):
        canvas.saveState(); canvas.setFont(PDF_FONT,8); canvas.setFillColor(colors.HexColor("#"+GRAY))
        canvas.drawCentredString(letter[0]/2,.3*inch,f"COINVIERTE · Ficha institucional · Página {_doc.page}"); canvas.restoreState()
    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return output.getvalue()
